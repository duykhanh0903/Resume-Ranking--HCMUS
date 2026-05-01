import os
import pdfplumber
import re
from docx import Document


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(BASE_DIR, '..'))

_TESSERACT_PATH = os.path.join(PROJECT_ROOT, 'vendor', 'Tesseract-OCR', 'tesseract.exe')

try:
    import pytesseract
    if os.path.exists(_TESSERACT_PATH):
        pytesseract.pytesseract.tesseract_cmd = _TESSERACT_PATH
    _TESSERACT_AVAILABLE = True
except ImportError:
    _TESSERACT_AVAILABLE = False


class ResumeBlockParser:
    def __init__(self, line_gap=4, block_gap=12):
        self.line_gap  = line_gap
        self.block_gap = block_gap

    def _is_scanned(self, page):
        try:
            text = page.extract_text() or ""
            return len(text.strip()) < 50
        except Exception:
            return False

    def _process_scanned_page(self, page):
        if not _TESSERACT_AVAILABLE:
            return []
        try:
            pix  = page.to_image(resolution=300).original
            text = pytesseract.image_to_string(pix, lang='eng')
            page_data = []
            for s in text.split('\n\n'):
                s = s.strip()
                if s:
                    page_data.append({"type": "FULL", "text": s.replace('\n', ' ')})
            return page_data
        except Exception as e:
            print(f"OCR failed: {e}")
            return []

    def _process_digital_page(self, page):
        words = page.extract_words(x_tolerance=3, y_tolerance=3)
        if not words:
            return []

        lines = []
        words.sort(key=lambda w: (w['top'], w['x0']))
        current_line = [words[0]]
        for i in range(1, len(words)):
            center_y_curr = (words[i]['top']         + words[i]['bottom'])         / 2
            center_y_prev = (current_line[-1]['top'] + current_line[-1]['bottom']) / 2
            if abs(center_y_curr - center_y_prev) <= self.line_gap:
                current_line.append(words[i])
            else:
                current_line.sort(key=lambda w: w['x0'])
                lines.append(current_line)
                current_line = [words[i]]
        current_line.sort(key=lambda w: w['x0'])
        lines.append(current_line)

        blocks_of_lines = []
        current_block = [lines[0]]
        for i in range(1, len(lines)):
            gap = lines[i][0]['top'] - current_block[-1][0]['bottom']
            if gap < self.block_gap:
                current_block.append(lines[i])
            else:
                blocks_of_lines.append(current_block)
                current_block = [lines[i]]
        blocks_of_lines.append(current_block)

        page_data = []
        for block in blocks_of_lines:
            block_words = [w for line in block for w in line]
            split_x     = self._find_gutter_in_block(block_words, page.width)
            if split_x:
                left  = sorted([w for w in block_words if w['x1'] <= split_x],  key=lambda w: (w['top'], w['x0']))
                right = sorted([w for w in block_words if w['x0'] >  split_x], key=lambda w: (w['top'], w['x0']))
                page_data.append({"type": "LEFT",  "text": " ".join(w['text'] for w in left)})
                page_data.append({"type": "RIGHT", "text": " ".join(w['text'] for w in right)})
            else:
                block_words.sort(key=lambda w: (w['top'], w['x0']))
                page_data.append({"type": "FULL", "text": " ".join(w['text'] for w in block_words)})
        return page_data

    def _find_gutter_in_block(self, words, page_width):
        width = int(page_width)
        occ = [0] * width
        for w in words:
            for x in range(int(w['x0']), int(w['x1'])):
                if 0 <= x < width:
                    occ[x] = 1
        max_gap, best_split, curr_gap = 0, None, 0
        for x in range(int(width * 0.2), int(width * 0.7)):
            if occ[x] == 0:
                curr_gap += 1
            else:
                if curr_gap > max_gap:
                    max_gap   = curr_gap
                    best_split = x - (curr_gap // 2)
                curr_gap = 0
        return best_split if max_gap > 30 else None

    def parse_file(self, file_path) -> tuple[list, list]:
        """Returns (blocks, embedded_links)"""
        ext = file_path.split('.')[-1].lower()
        if ext == 'pdf':
            return self.parse_pdf(file_path)
        elif ext == 'docx':
            return self.parse_docx(file_path), self._extract_docx_links(file_path)
        else:
            print(f"Unsupported format: {ext}")
            return [], []

    def parse_pdf(self, pdf_path) -> tuple[list, list]:
        """Returns (blocks, embedded_links)"""
        all_extracted_blocks = []
        all_links = []

        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                try:
                    # Collect embedded hyperlinks from annotations
                    all_links.extend(self._extract_hyperlinks(page))

                    if self._is_scanned(page):
                        page_blocks = self._process_scanned_page(page)
                    else:
                        page_blocks = self._process_digital_page(page)
                    if page_blocks:
                        all_extracted_blocks.extend(page_blocks)
                except Exception as e:
                    print(f"Skipping page: {e}")
                    continue

        return all_extracted_blocks, list(set(all_links))

    def parse_docx(self, docx_path):
        doc = Document(docx_path)
        page_data = []
        for element in doc.element.body:
            if element.tag.endswith('p'):
                text = "".join(
                    node.text for node in element.findall(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
                    )
                )
                if text.strip():
                    page_data.append({"type": "FULL", "text": text.strip()})
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        table_text = self._extract_table_text(table)
                        if table_text:
                            page_data.append({"type": "FULL", "text": table_text})
        return page_data
    
    def _extract_docx_links(self, docx_path) -> list[str]:
        """Extract hyperlinks from DOCX relationships."""
        links = []
        try:
            from docx import Document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            doc = Document(docx_path)
            for rel in doc.part.rels.values():
                if "hyperlink" in rel.reltype:
                    url = rel._target
                    if isinstance(url, str) and url.startswith("http"):
                        links.append(url.strip())
        except Exception:
            pass
        return list(set(links))

    def _extract_hyperlinks(self, page) -> list[str]:
        links = []
        try:
            if page.annots:
                for annot in page.annots:
                    # pdfplumber exposes annotation dict
                    uri = annot.get("uri") or annot.get("URI")
                    if uri and isinstance(uri, str) and uri.startswith("http"):
                        links.append(uri.strip())
        except Exception:
            pass
        return links

    def _extract_table_text(self, table):
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                if not cells or cell.text != cells[-1]:
                    cells.append(cell.text.strip())
            row_content = " | ".join(filter(None, cells))
            if row_content:
                rows.append(row_content)
        return "\n".join(rows)

    def restructure_data(self, all_blocks):
        if not all_blocks:
            return ""
        final_fragments = []
        for block in all_blocks:
            text = block.get('text', '').strip()
            if not text:
                continue
            text = re.sub(r'[|\\_~*•►«»©®™¥%]', '', text)
            text = re.sub(r'(\w+)-\s+(\w+)', r'\1\2', text)
            text = re.sub(r'\s+', ' ', text)
            final_fragments.append(text)
        return "\n\n".join(final_fragments).strip()